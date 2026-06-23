from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "fixed_final_report_data_security.docx"
OUT = ROOT / "outputs" / "数据分类分级保护系统-课程设计报告-优化版.docx"
SHOT_DIR = ROOT / "outputs" / "project_screenshots"


REPLACEMENTS = {
    "基于分类分级的数据安全治理系统": "数据分类分级保护系统",
    "设计包含 20 张核心表的概念模型和逻辑模型，并梳理实体关系、主外键和约束。": (
        "设计包含 16 张业务核心表和 4 张 Agent 扩展表的概念模型与逻辑模型，并梳理实体关系、主外键和约束。"
    ),
    "系统采用前后端分离架构。前端使用 Vue 3、Vite、Element Plus 和 Axios，负责登录、资产列表、分类分级、审批、审计、看板和 Agent 工作台等交互界面；后端使用 Spring Boot 3、Spring Web、Spring JDBC 和 Spring Security，对外提供 REST API；数据库层使用 SQL Server 作为主设计目标，同时使用 H2 内存库便于课程演示和快速启动。": (
        "系统采用前后端分离架构。前端使用 Vue 3、Vite、Element Plus、Axios 和 Vue Router，负责登录、资产管理、分类分级、审批、审计、看板、Agent 工作台和悬浮助手等交互界面；后端使用 Java 17、Spring Boot 3、Spring Web、Spring JDBC、Spring Security 和 LangChain4j，对外提供 REST API 与 OpenAI-compatible 模型编排；数据库层使用 MySQL，启动时自动执行 schema.sql 与 data.sql 初始化演示数据。"
    ),
    "物理模型以 SQL Server 为主，使用 bigint identity 作为主键，nvarchar 存储中文业务字段，datetime2 存储时间，decimal(5,2) 存储置信度。针对查询高频字段建立索引，包括字段名、分类等级、审计用户、审计时间、申请状态、Agent 任务类型、推荐目标、风险等级和问答用户等。": (
        "物理模型以 MySQL 为主，使用 bigint auto_increment 作为主键，varchar/text 存储中文业务字段，timestamp 存储时间，decimal(5,2) 存储置信度。针对查询高频字段建立索引，包括字段名、分类等级、审计用户、审计时间、申请状态、Agent 任务类型、推荐目标、风险等级和问答用户等。"
    ),
    "Agent 模块包含字段分类、审批建议、审计风险、安全报告和问答服务。它综合字段名、样例值、分类规则、申请理由和审计日志输出建议、置信度、告警和整改方案。": (
        "Agent 模块包含字段分类、审批建议、审计风险、安全报告和问答服务。它综合字段名、样例值、分类规则、申请理由和审计日志输出建议、置信度、告警和整改方案；默认提供本地规则推理，外部模型通道可通过 DeepSeek 或硅基流动的 OpenAI-compatible 接口接入，且 API Key 仅由后端环境变量读取。"
    ),
    "由于本报告由源码和脚本生成，系统截图可在正式汇报前按实际运行界面补充到本节。当前正文已覆盖测试环境、测试步骤、预期结果与结果分析。": (
        "本次优化已按当前运行界面补充截图，覆盖首页看板、Agent 工作台、安全报告和悬浮助手，便于答辩时直接对应系统闭环。当前正文已覆盖测试环境、测试步骤、预期结果与结果分析。"
    ),
    "设计 20 张核心表，覆盖用户权限、数据资产、分类分级、访问审批、审计日志和 Agent 治理。": (
        "设计 16 张业务核心表和 4 张 Agent 扩展表，覆盖用户权限、数据资产、分类分级、访问审批、审计日志和 Agent 治理。"
    ),
    "测试环境使用 Windows 本地开发环境，后端通过 Maven 启动 Spring Boot 服务，默认连接 H2 内存数据库并加载 schema.sql 与 data.sql；前端通过 npm run dev 启动 Vite 服务；浏览器访问 http://localhost:5173。系统也提供 SQL Server 初始化脚本，可执行 sql/sqlserver_schema_init.sql 完成数据库对象创建。": (
        "测试环境使用 Windows 本地开发环境，后端通过 Maven 启动 Spring Boot 服务，默认连接 MySQL 数据库并加载 schema.sql 与 data.sql；前端通过 npm run dev 启动 Vite 服务；浏览器访问 http://localhost:5173。系统也提供 sql/mysql_schema_init.sql 初始化脚本，便于课程演示前快速创建数据库对象和演示数据。"
    ),
    "测试环境使用 Windows 本地开发环境，后端通过 Maven 启动 Spring Boot 服务，默认连接 H2 内存数据库并加载 schema.sql 与 data.sql；前端通过 npm run dev 启动 Vite 服务；浏览器访问 http://localhost:5173。系统也提供 SQL Server 初始化脚本，可执行 sql/sqlserver_schema_init.sql 后切换 sqlserver profile。": (
        "测试环境使用 Windows 本地开发环境，后端通过 Maven 启动 Spring Boot 服务，默认连接 MySQL 数据库并加载 schema.sql 与 data.sql；前端通过 npm run dev 启动 Vite 服务；浏览器访问 http://localhost:5173。系统也提供 sql/mysql_schema_init.sql 初始化脚本，便于课程演示前快速创建数据库对象和演示数据。"
    ),
    "编写 SQL Server 建库脚本，实现表、索引、视图、存储过程和触发器。": (
        "编写 MySQL 建库与初始化脚本，实现表、索引、视图、存储过程和触发器。"
    ),
    "完成 SQL Server 建库脚本，包含约束、索引、视图、存储过程、触发器和演示数据。": (
        "完成 MySQL 建库与初始化脚本，包含约束、索引、视图、存储过程、触发器和演示数据。"
    ),
    "系统截图、录屏和真实小组信息需要在正式提交前补充完善。": (
        "真实小组信息和必要的答辩录屏仍需在正式提交前补充；系统截图已按当前运行界面更新到正文。"
    ),
    "D:/database/backend，D:/database/frontend，D:/database/sql/sqlserver_schema_init.sql": (
        "E:/Desktop/Safedatabase/backend，E:/Desktop/Safedatabase/frontend，E:/Desktop/Safedatabase/sql/mysql_schema_init.sql"
    ),
}


FIGURES = [
    (
        "01_dashboard.png",
        "图 4-1 首页看板：集中展示数据源、数据表、字段资产、敏感字段、待审批申请、风险告警和分类覆盖率。",
    ),
    (
        "03_agent_workbench.png",
        "图 4-2 Agent 工作台：汇总字段分类、审批建议、风险告警、安全报告等治理任务入口。",
    ),
    (
        "07_security_report_generated.png",
        "图 4-3 安全报告：基于当前数据生成资产概览、分类覆盖率、风险情况和整改建议。",
    ),
    (
        "06_floating_ai.png",
        "图 4-4 悬浮 AI 助手：在任意页面提供本地规则推理和问答入口，不改变原有权限边界。",
    ),
]


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def replace_text(paragraph):
    text = paragraph.text
    new_text = text
    for old, new in REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    if new_text == text:
        return
    paragraph.clear()
    run = paragraph.add_run(new_text)
    set_run_font(run)


def insert_picture_before(target, image_path, caption):
    pic_para = target.insert_paragraph_before()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))

    cap_para = target.insert_paragraph_before()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(8)
    cap_run = cap_para.add_run(caption)
    set_run_font(cap_run, size=9.5, color=RGBColor(90, 90, 90))


def insert_page_break_before(target):
    paragraph = target.insert_paragraph_before()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    for paragraph in doc.paragraphs:
        replace_text(paragraph)

    heading5 = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5 总结与展望" and paragraph.style.name == "Heading 1":
            heading5 = paragraph
            break
    if heading5 is None:
        raise RuntimeError("未找到“5 总结与展望”插入位置")

    insert_page_break_before(heading5)
    h = heading5.insert_paragraph_before("4.4 运行截图与效果对齐")
    h.style = "Heading 2"
    for run in h.runs:
        set_run_font(run, size=13, bold=True, color=RGBColor(46, 116, 181))

    intro = heading5.insert_paragraph_before(
        "以下截图来自当前已启动的项目运行界面，补充说明系统已从文档设计落到可演示页面。截图按“治理态势-智能辅助-报告整改-全局问答”的顺序组织，与答辩 PPT 的演示流程保持一致。"
    )
    intro.paragraph_format.first_line_indent = Pt(22)
    intro.paragraph_format.line_spacing = 1.25
    intro.paragraph_format.space_after = Pt(6)
    for run in intro.runs:
        set_run_font(run)

    for filename, caption in FIGURES:
        image_path = SHOT_DIR / filename
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        insert_picture_before(heading5, image_path, caption)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
