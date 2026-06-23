# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("D:/database")
OUT = ROOT / "final_report_data_security_utf8.docx"
ASSETS = ROOT / "tmp" / "report_assets_utf8"
ASSETS.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


def get_font(size=22, bold=False):
    names = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for name in names:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill="#EEF4FB", outline="#2E74B5"):
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    font = get_font(22, True)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = len(lines) * 30
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        box = draw.textbbox((0, 0), line, font=font)
        w = box[2] - box[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill="#0B2545")
        y += 30


def draw_arrow(draw, a, b, color="#44546A"):
    import math

    draw.line([a, b], fill=color, width=3)
    angle = math.atan2(b[1] - a[1], b[0] - a[0])
    for delta in (2.6, -2.6):
        p = (b[0] - 14 * math.cos(angle + delta), b[1] - 14 * math.sin(angle + delta))
        draw.line([b, p], fill=color, width=3)


def make_architecture():
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((650, 28), "系统总体架构图", font=get_font(36, True), fill="#0B2545")
    draw_box(d, (80, 150, 310, 250), "Vue 3 前端\nElement Plus")
    draw_box(d, (440, 130, 700, 270), "Spring Boot\nREST API")
    modules = [
        ("认证与角色", 820, 80),
        ("资产管理", 820, 185),
        ("分类分级", 820, 290),
        ("访问审批", 820, 395),
        ("审计日志", 820, 500),
        ("治理 Agent", 820, 605),
    ]
    for text, x, y in modules:
        draw_box(d, (x, y, x + 230, y + 75), text, fill="#F7FBFF")
    draw_box(d, (1150, 160, 1490, 330), "MySQL\n表、索引\n初始化数据", fill="#FFF8E8", outline="#C78500")
    draw_box(d, (1150, 480, 1490, 665), "Agent 子服务\n字段分类 / 审批建议\n风险分析 / 安全报告 / 问答", fill="#F4F6F9", outline="#6B7280")
    draw_arrow(d, (310, 200), (440, 200))
    for _, x, y in modules:
        draw_arrow(d, (700, 200), (x, y + 37))
    for _, x, y in modules[:5]:
        draw_arrow(d, (1050, y + 37), (1150, 245))
    draw_arrow(d, (1050, 642), (1150, 570))
    draw_arrow(d, (1320, 480), (1320, 330))
    path = ASSETS / "architecture.png"
    img.save(path)
    return path


def make_workflow():
    img = Image.new("RGB", (1600, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((650, 28), "核心业务流程图", font=get_font(36, True), fill="#0B2545")
    nodes = [
        ("登记数据源\n表与字段", 80, 130),
        ("自动/人工\n分类分级", 350, 130),
        ("按等级判断\n访问方式", 620, 130),
        ("L1-L3\n脱敏或授权展示", 910, 90),
        ("L4-L5\n提交访问申请", 910, 220),
        ("审批人结合\nAgent 建议处理", 1200, 220),
        ("查看原始值\n或驳回", 1200, 380),
        ("写入审计日志", 750, 540),
        ("风险分析\n生成告警", 1030, 540),
        ("安全报告\n整改闭环", 1290, 540),
    ]
    for text, x, y in nodes:
        draw_box(d, (x, y, x + 230, y + 90), text)
    for a, b in [
        ((310, 175), (350, 175)),
        ((580, 175), (620, 175)),
        ((850, 155), (910, 135)),
        ((850, 195), (910, 265)),
        ((1140, 265), (1200, 265)),
        ((1315, 310), (1315, 380)),
        ((1025, 180), (865, 540)),
        ((1315, 470), (865, 540)),
        ((980, 585), (1030, 585)),
        ((1260, 585), (1290, 585)),
    ]:
        draw_arrow(d, a, b)
    path = ASSETS / "workflow.png"
    img.save(path)
    return path


def make_er():
    img = Image.new("RGB", (1700, 1100), "white")
    d = ImageDraw.Draw(img)
    d.text((735, 28), "E-R 核心关系图", font=get_font(36, True), fill="#0B2545")
    nodes = {
        "用户": (80, 110),
        "角色": (350, 110),
        "权限": (620, 110),
        "数据源": (80, 330),
        "数据表资产": (360, 330),
        "字段资产": (650, 330),
        "分类类别": (980, 240),
        "安全等级": (980, 400),
        "字段分类结果": (1260, 330),
        "分类规则": (1260, 160),
        "访问申请": (980, 590),
        "审批记录": (1260, 590),
        "审计日志": (650, 590),
        "Agent任务": (80, 820),
        "Agent建议": (360, 820),
        "风险告警": (650, 820),
        "问答历史": (980, 820),
    }
    for text, (x, y) in nodes.items():
        draw_box(d, (x, y, x + 230, y + 80), text, fill="#F7FBFF")
    edges = [
        ("用户", "角色"),
        ("角色", "权限"),
        ("数据源", "数据表资产"),
        ("数据表资产", "字段资产"),
        ("字段资产", "字段分类结果"),
        ("分类类别", "字段分类结果"),
        ("安全等级", "字段分类结果"),
        ("分类类别", "分类规则"),
        ("安全等级", "分类规则"),
        ("字段资产", "访问申请"),
        ("用户", "访问申请"),
        ("访问申请", "审批记录"),
        ("用户", "审计日志"),
        ("用户", "Agent任务"),
        ("字段资产", "Agent建议"),
        ("访问申请", "Agent建议"),
        ("审计日志", "风险告警"),
        ("用户", "问答历史"),
    ]
    for a, b in edges:
        ax, ay = nodes[a]
        bx, by = nodes[b]
        draw_arrow(d, (ax + 230, ay + 40), (bx, by + 40))
    path = ASSETS / "er.png"
    img.save(path)
    return path


def set_run(run, name="宋体", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", first=True, align=None, after=6, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(after)
    if first:
        para.paragraph_format.first_line_indent = Pt(22)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        set_run(run, size=size)
    return para


def add_heading(doc, text, level=1):
    para = doc.add_paragraph(style=f"Heading {level}")
    run = para.add_run(text)
    set_run(run, "黑体", {1: 16, 2: 13, 3: 12}.get(level, 11), True, BLUE if level < 3 else DARK_BLUE)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run(run, size=10.5)


def add_number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run(run, size=10.5)


def shade_cell(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, center=False, size=9.5):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    set_run(run, size=size, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, True, True)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=9 if len(str(value)) > 55 else 9.5)
    add_para(doc, "", False, after=2)
    return table


def add_caption(doc, text):
    para = add_para(doc, text, False, WD_ALIGN_PARAGRAPH.CENTER, 6, 9.5)
    for run in para.runs:
        run.font.color.rgb = GRAY


def configure_doc(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE if name != "Heading 3" else DARK_BLUE
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build():
    arch = make_architecture()
    workflow = make_workflow()
    er = make_er()

    doc = Document()
    configure_doc(doc)

    para = add_para(doc, "数据库系统及安全", False, WD_ALIGN_PARAGRAPH.CENTER, 8, 18)
    set_run(para.runs[0], "黑体", 18, True)
    para = add_para(doc, "课程设计报告", False, WD_ALIGN_PARAGRAPH.CENTER, 50, 24)
    set_run(para.runs[0], "黑体", 24, True)

    cover_rows = [
        ("课 程 号：", "（请填写）"),
        ("课 序 号：", "（请填写）"),
        ("课 程 班：", "（请填写）"),
        ("项 目 名：", "基于分类分级的数据安全治理系统"),
        ("组    号：", "（请填写）"),
        ("项目组长：", "（姓名 学号 请填写）"),
        ("项目组员：", "（姓名 学号 请填写）"),
        ("汇 报 人：", "（姓名 学号 请填写）"),
        ("指导老师：", "龚勋"),
        ("评阅成绩：", ""),
        ("评阅意见：", ""),
    ]
    for label, value in cover_rows:
        para = add_para(doc, "", False, after=8, size=14)
        para.paragraph_format.left_indent = Inches(1.35)
        run = para.add_run(label)
        set_run(run, size=14, bold=True)
        run = para.add_run("    " + value)
        set_run(run, size=14)
    add_para(doc, "提交报告时间：2026 年 6 月 23 日", False, WD_ALIGN_PARAGRAPH.CENTER, 10, 12)
    add_para(doc, "四川大学网络空间安全学院", False, WD_ALIGN_PARAGRAPH.CENTER, 0, 14)
    doc.add_page_break()

    add_heading(doc, "工作量划分表")
    add_table(
        doc,
        ["成员姓名", "成员学号", "工作内容", "工作量"],
        [
            ["组长（请填写）", "（请填写）", "总体协调；需求分析；数据库总体架构、概念模型与报告统稿", "35%"],
            ["成员A（请填写）", "（请填写）", "后端接口、认证授权、访问审批、审计日志与安全控制实现", "25%"],
            ["成员B（请填写）", "（请填写）", "前端页面、数据资产管理、分类分级、脱敏展示与效果演示", "25%"],
            ["成员C（请填写）", "（请填写）", "SQL 脚本、测试用例、文档整理、图表与汇报材料", "15%"],
        ],
    )
    add_para(doc, "说明：正式提交前请替换为小组真实姓名、学号和实际工作量。", False, size=10.5)
    doc.add_page_break()

    add_heading(doc, "摘要")
    add_para(
        doc,
        "本项目面向企业数据资产管理和数据安全治理场景，设计并实现“基于分类分级的数据安全治理系统”。系统围绕数据源、数据表、字段资产、分类类别、安全等级、分类规则、脱敏策略、访问申请、审批记录、审计日志和治理 Agent 等核心对象建立统一数据库模型，支持字段资产登记、敏感字段识别、L1-L5 分级保护、规则自动分类、动态脱敏展示、L4/L5 高敏字段访问审批、审计留痕、风险告警和安全报告生成。数据库设计遵循实体完整性、参照完整性、最小权限、可审计和可追溯原则，完成概念模型、逻辑模型、物理模型与 BCNF 规范化分析，并在 MySQL 初始化脚本中实现核心表、唯一约束、索引和演示数据。系统实现采用 Spring Boot 3、Spring JDBC、Spring Security、Vue 3、Element Plus 和 MySQL，形成前后端分离架构。项目还引入本地规则型数据安全治理 Agent，实现字段智能分类、访问审批建议、审计风险分析、Markdown 安全报告和问答辅助，构建“识别-防护-审批-审计-告警-整改”的闭环。测试结果表明，系统能够完整支撑数据分类分级保护流程，具备较好的安全性、可扩展性和课程演示价值。",
    )
    add_para(doc, "主题词：数据安全；数据库设计；分类分级；动态脱敏；访问审批；审计日志；治理 Agent", False)
    add_heading(doc, "Abstract")
    add_para(
        doc,
        "This project designs and implements a data security governance system based on data classification and grading. It builds a unified database model around data assets, field-level classification, masking policies, access requests, approval records, audit logs, and governance agents. The system supports asset registration, sensitive-field identification, L1-L5 grading, rule-based automatic classification, dynamic masking, approval control for L4/L5 raw-value access, audit tracing, risk alerts, and security report generation. The database design follows integrity, least privilege, auditability, and traceability principles, and implements core tables, constraints, indexes, and demo data in MySQL. The implementation adopts Spring Boot 3, Spring JDBC, Spring Security, Vue 3, Element Plus, and MySQL. A local rule-based governance agent is added to provide classification suggestions, access review suggestions, audit-risk analysis, security reports, and question-answer assistance.",
        False,
    )
    add_para(doc, "Keywords: Data Security; Database Design; Classification and Grading; Dynamic Masking; Access Approval; Audit Log; Governance Agent", False)
    doc.add_page_break()

    add_heading(doc, "目录")
    for item in [
        "1 前言",
        "1.1 选题背景与意义",
        "1.2 国内外研究现状",
        "1.3 本项目的主要工作",
        "2 系统设计",
        "2.1 系统总体设计",
        "2.2 数据库设计",
        "3 系统实现",
        "4 系统测试",
        "5 总结与展望",
        "参考文献",
    ]:
        add_para(doc, item, False, after=2)
    doc.add_page_break()

    add_heading(doc, "1 前言")
    add_heading(doc, "1.1 选题背景与意义", 2)
    add_para(doc, "在数字化业务系统中，数据库已经从单纯的数据存储设施演变为业务运行、权限控制、合规审计和风险治理的基础平台。客户资料、交易订单、账号凭证、运维日志等数据一旦缺少分类分级管理，系统很难判断哪些字段可以公开展示、哪些字段必须脱敏、哪些字段需要审批后才能访问。《网络安全法》《数据安全法》《个人信息保护法》以及数据分类分级相关标准均强调对重要数据和个人信息进行识别、分级、访问控制、审计和安全评估。")
    add_para(doc, "本项目把数据库系统设计与网络空间安全要求结合起来：一方面训练需求分析、E-R 建模、关系模式转换、规范化、SQL 约束、索引、视图、存储过程和触发器等数据库设计能力；另一方面围绕敏感数据识别、动态脱敏、访问审批、审计留痕和风险告警实现可演示的安全闭环。")
    add_heading(doc, "1.2 国内外研究现状", 2)
    add_para(doc, "当前数据安全治理系统通常包括数据资产发现、元数据管理、敏感数据识别、数据分类分级、权限审批、脱敏加密、审计分析和合规报告等能力。成熟产品强调跨源扫描、策略编排和自动化治理；课程设计系统则需要用清晰、可验证的数据库结构表达安全治理逻辑。本项目从字段级资产入手，将分类规则、分级结果、脱敏策略、审批记录和审计日志全部落库，并加入本地规则型 Agent，体现智能化治理趋势。")
    add_heading(doc, "1.3 本项目的主要工作", 2)
    for item in [
        "完成数据资产、分类分级、脱敏策略、访问审批、审计日志和 Agent 治理需求分析。",
        "设计包含 20 张核心表的概念模型和逻辑模型，并梳理实体关系、主外键和约束。",
        "完成关系模式规范化分析，说明主要关系满足 BCNF 的依据。",
        "编写 MySQL 初始化脚本，实现核心表、索引和演示数据。",
        "基于 Spring Boot 与 Vue 实现前后端分离系统，覆盖登录、资产管理、分类分级、脱敏、审批、审计和 Agent 页面。",
        "完成测试用例设计与结果分析，验证核心业务流程和数据库安全控制。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2 系统设计")
    add_heading(doc, "2.1 系统总体设计", 2)
    add_heading(doc, "2.1.1 系统架构设计", 3)
    add_para(doc, "系统采用前后端分离架构。前端使用 Vue 3、Vite、Element Plus 和 Axios，负责登录、资产列表、分类分级、审批、审计、看板和 Agent 工作台等交互界面；后端使用 Spring Boot 3、Spring Web、Spring JDBC 和 Spring Security，对外提供 REST API；数据库层使用 MySQL 作为系统运行数据库，启动时加载 schema.sql 与 data.sql 完成演示数据初始化。")
    doc.add_picture(str(arch), width=Inches(6.5))
    add_caption(doc, "图2-1 系统总体架构图")
    add_heading(doc, "2.1.2 核心业务流程设计", 3)
    add_para(doc, "核心流程以字段资产为治理粒度：安全管理员先登记数据源、数据表和字段，再通过分类规则或 Agent 建议生成字段分类分级结果。当普通用户访问 L4/L5 高敏字段原始值时，系统要求提交访问申请，由审批人员结合业务理由和 Agent 建议进行审批。所有查看、申请、审批、分类修改和 Agent 操作都写入审计日志，风险分析服务再从日志中识别异常访问并生成风险告警。")
    doc.add_picture(str(workflow), width=Inches(6.5))
    add_caption(doc, "图2-2 核心业务流程图")
    add_heading(doc, "2.1.3 创新点与安全性考虑", 3)
    for item in [
        "字段级分类分级：以 data_field_asset 为最小治理对象，适合对字段、约束、关系和查询进行细粒度表达。",
        "动态脱敏与审批联动：L3 数据默认脱敏展示，L4/L5 原始值需要有效审批记录，体现最小权限原则。",
        "全链路审计：登录、字段查看、原始值访问、分类更新、申请审批、Agent 操作均写入 audit_log。",
        "数据库对象完整：除基本表外，还设计索引、视图、存储过程和触发器，体现数据库系统综合能力。",
        "本地规则 Agent：在不依赖外部模型的情况下完成分类建议、审批建议、风险告警和安全报告。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 数据库设计", 2)
    add_heading(doc, "2.2.1 概念模型设计", 3)
    add_para(doc, "概念模型围绕“用户-角色-权限”“数据资产-字段-分类分级”“访问申请-审批-审计”“Agent 任务-建议-告警-问答”四组实体展开。用户通过角色获得权限；数据源包含多张数据表，数据表包含多个字段；字段最多对应一条当前分类分级结果；分类规则指向分类类别和安全等级；普通用户访问高敏字段时产生访问申请，审批动作写入审批记录；系统关键操作写入审计日志；Agent 模块基于字段、申请和日志生成任务、建议、告警和问答历史。")
    doc.add_picture(str(er), width=Inches(6.5))
    add_caption(doc, "图2-3 E-R 核心关系图")
    add_heading(doc, "2.2.1.1 实体识别与属性定义", 3)
    add_table(
        doc,
        ["实体", "主要属性", "设计说明"],
        [
            ["sys_user", "id, username, password, real_name, email, phone, enabled", "系统用户，参与登录、申请、审批和审计。"],
            ["sys_role / sys_permission", "role_code, permission_code, description", "角色和权限字典，通过关联表实现 RBAC。"],
            ["data_source / data_table_asset", "source_name, table_name, business_name, owner_department", "登记数据源与数据表资产。"],
            ["data_field_asset", "table_id, field_name, field_type, sample_value, is_sensitive", "字段级资产，是分类分级和访问控制的核心粒度。"],
            ["classification_category / classification_level", "category_name, level_code, level_order", "定义数据类别和 L1-L5 安全等级。"],
            ["classification_rule", "match_type, match_pattern, category_id, level_id", "自动识别规则，支持 keyword 和 regex。"],
            ["field_classification", "field_id, category_id, level_id, classify_method", "字段当前分类分级结果。"],
            ["masking_policy", "policy_type, example_before, example_after", "动态脱敏策略样例。"],
            ["access_request / approval_record", "user_id, field_id, status, approval_result", "高敏数据访问申请和审批记录。"],
            ["audit_log", "operation_type, target_type, operation_time, result", "关键操作审计与追溯。"],
            ["Agent 表", "task_type, confidence, risk_level, question, answer", "治理 Agent 的任务、建议、告警和问答存储。"],
        ],
    )
    add_heading(doc, "2.2.1.2 实体间关系分析", 3)
    for item in [
        "data_source 与 data_table_asset 为 1:N；data_table_asset 与 data_field_asset 为 1:N。",
        "data_field_asset 与 field_classification 为 1:0..1，保证一个字段只有一条当前分类结果。",
        "classification_category、classification_level 分别与 field_classification、classification_rule 为 1:N。",
        "sys_user、sys_role、sys_permission 通过 sys_user_role、sys_role_permission 实现多对多。",
        "access_request 指向申请用户和目标字段，approval_record 指向访问申请和审批人。",
        "audit_log 可关联用户与目标对象，为风险分析提供事实依据。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2.2 逻辑模型设计", 3)
    add_heading(doc, "2.2.2.1 关系模式转换", 3)
    add_para(doc, "根据 E-R 模型，将强实体转换为独立关系，将 M:N 关系转换为关联关系，并将 1:N 关系通过外键表达。字段级分类结果采用 field_id 唯一约束控制“一个字段最多一条当前分类结果”；角色权限采用联合主键控制重复授权；访问审批通过 access_request 与 approval_record 分离，既保留申请状态，又保留每次审批动作。")
    add_table(
        doc,
        ["关系模式", "主键", "外键/约束"],
        [
            ["sys_user(id, username, password, real_name, email, phone, enabled)", "id", "username unique"],
            ["sys_user_role(user_id, role_id)", "(user_id, role_id)", "user_id -> sys_user, role_id -> sys_role"],
            ["data_table_asset(id, source_id, table_name, business_name)", "id", "source_id -> data_source"],
            ["data_field_asset(id, table_id, field_name, field_type, sample_value)", "id", "table_id -> data_table_asset"],
            ["field_classification(id, field_id, category_id, level_id, classify_method)", "id", "field_id unique; category_id -> category; level_id -> level"],
            ["access_request(id, user_id, field_id, reason, status, valid_until)", "id", "status in PENDING/APPROVED/REJECTED"],
            ["approval_record(id, request_id, approver_id, approval_result)", "id", "request_id -> access_request; approver_id -> sys_user"],
            ["audit_log(id, user_id, operation_type, target_type, target_id, operation_time)", "id", "operation_type not null"],
        ],
    )
    add_heading(doc, "2.2.2.2 规范化", 3)
    add_para(doc, "规范化目标是消除插入、删除和更新异常，使每个关系中的非主属性完全依赖于候选键，并避免传递依赖。本项目的关系模式均以单一业务主题建表，字典类信息独立存放，关联关系独立建表，因此主要关系可以达到 BCNF。")
    for item in [
        "sys_user 中 username 是候选键，real_name、email、phone、enabled 等属性依赖于 id 或 username。",
        "data_field_asset 中字段属性依赖于字段 id，表信息通过 table_id 外键引用，不重复保存数据源或部门信息。",
        "field_classification 中 field_id 具有唯一性，category_id、level_id、classify_method、classified_time 都依赖于 field_id 或 id。",
        "classification_rule 中规则内容、命中类型、目标分类和目标等级依赖于规则 id；分类名称、等级名称由字典表维护。",
        "access_request 与 approval_record 分离后，申请状态依赖于申请 id，审批意见依赖于审批记录 id。",
    ]:
        add_number(doc, item)
    add_heading(doc, "2.2.2.3 关系模式清单", 3)
    add_table(
        doc,
        ["类别", "表名", "作用"],
        [
            ["用户权限", "sys_user, sys_role, sys_permission, sys_user_role, sys_role_permission", "实现用户、角色、权限和授权关系。"],
            ["数据资产", "data_source, data_table_asset, data_field_asset", "实现数据源、表、字段三级资产目录。"],
            ["分类分级", "classification_category, classification_level, classification_rule, field_classification", "实现类别、等级、规则和字段结果。"],
            ["访问安全", "masking_policy, access_request, approval_record, audit_log", "实现脱敏、申请、审批和审计。"],
            ["Agent治理", "agent_task, agent_recommendation, risk_alert, agent_chat_history", "实现智能任务、建议、风险告警和问答历史。"],
        ],
    )
    add_heading(doc, "2.2.3 物理模型设计", 3)
    add_para(doc, "物理模型以 MySQL 为主，使用 bigint auto_increment 作为主键，varchar/text 存储业务字段，timestamp 存储时间，decimal(5,2) 存储置信度。针对查询高频字段建立索引，包括字段名、分类等级、审计用户、审计时间、申请状态、Agent 任务类型、推荐目标、风险等级和问答用户等。")
    add_table(
        doc,
        ["物理对象", "名称", "用途"],
        [
            ["索引", "idx_data_field_name", "提高字段名查询和规则匹配定位效率。"],
            ["索引", "idx_field_classification_level", "提高按 L1-L5 等级统计和筛选效率。"],
            ["索引", "idx_audit_user / idx_audit_time", "支撑审计日志按用户、时间窗口分析。"],
            ["索引", "idx_access_status", "支撑审批列表按状态查询。"],
            ["视图", "sensitive_field_view", "集中展示 L3 及以上敏感字段。"],
            ["视图", "user_accessible_field_view", "计算用户对字段原始值的可访问状态。"],
            ["存储过程", "sp_count_sensitive_fields_by_source", "按数据源统计各等级敏感字段数量。"],
            ["存储过程", "sp_auto_classify_fields", "根据启用规则批量维护字段分类结果。"],
            ["触发器", "trg_audit_field_classification_update", "分类分级更新后自动写入审计日志。"],
            ["触发器", "trg_audit_access_request_update", "访问申请状态更新后自动写入审计日志。"],
        ],
    )
    add_heading(doc, "2.2.4 数据库安全性设计", 3)
    for item in [
        "身份认证：系统提供登录接口，前端在请求头携带 Bearer token，后端识别当前用户。",
        "角色权限：设计 admin、security_admin、user、approver 四类角色，表达管理、分类、审批、审计等能力边界。",
        "分级访问：L1-L3 字段允许按规则展示，L4/L5 原始值必须存在有效 APPROVED 申请。",
        "动态脱敏：手机号、邮箱、身份证、银行卡、密码密钥等字段根据字段名和策略进行格式化脱敏展示。",
        "审计追溯：登录、字段详情查看、原始值查看、分类更新、访问申请、审批、Agent 操作均写入 audit_log。",
        "数据库约束：主键、唯一约束、检查约束和外键保证数据完整性，避免孤立审批、重复分类和非法状态。",
        "风险告警：Agent 通过高频访问、高敏字段集中访问、登录失败、同 IP 多账号、非管理员分类修改等规则生成风险告警。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.2.5 关键 SQL 脚本", 3)
    add_table(
        doc,
        ["脚本片段", "设计意图"],
        [
            ["field_classification.field_id unique", "保证一个字段只有一条当前分类分级结果，避免冲突等级。"],
            ["classification_rule.match_type check in ('keyword','regex')", "保证规则类型合法。"],
            ["access_request.status check in ('PENDING','APPROVED','REJECTED')", "保证审批状态闭环明确。"],
            ["sensitive_field_view where level_order >= 3", "集中筛选 L3 及以上敏感字段。"],
            ["user_accessible_field_view", "把字段等级、审批状态和有效期合并为可访问性判断。"],
            ["sp_auto_classify_fields", "用启用规则批量维护字段分类结果并同步 is_sensitive。"],
            ["trg_audit_access_request_update", "访问申请状态变化时自动写入审计日志。"],
        ],
    )

    add_heading(doc, "3 系统实现")
    add_heading(doc, "3.1 开发环境与工具", 2)
    add_table(
        doc,
        ["类别", "工具/版本", "说明"],
        [
            ["后端", "Java 17, Spring Boot 3.3.5, Spring Web, Spring JDBC, Spring Security, Lombok", "实现 REST API、认证上下文、数据库访问和服务层逻辑。"],
            ["数据库", "MySQL 8", "系统运行数据库，启动时自动初始化演示表和数据。"],
            ["前端", "Vue 3, Vite, Element Plus, Axios, Vue Router", "实现单页应用、管理界面和 Agent 页面。"],
            ["构建", "Maven, npm", "分别构建后端 JAR 和前端静态资源。"],
        ],
    )
    add_heading(doc, "3.2 模块实现", 2)
    modules = [
        ("3.2.1 登录与角色管理模块", "认证模块提供 /api/auth/login 与 /api/auth/me 接口。登录成功后生成 token，前端存入 localStorage 并在后续请求中携带 Authorization 请求头。"),
        ("3.2.2 数据资产管理模块", "数据资产模块提供数据源、数据表和字段资产的增删改查接口。字段详情接口还提供脱敏样例值和原始样例值访问逻辑；无有效审批访问 L4/L5 字段时只返回脱敏值并写入拒绝审计。"),
        ("3.2.3 分类分级与规则模块", "分类分级模块支持分类字典、等级字典、字段分类结果、分类规则维护和自动分类。自动分类根据字段名与启用规则的 match_pattern 进行匹配，命中后写入 field_classification。"),
        ("3.2.4 脱敏策略与访问审批模块", "MaskingService 根据字段名识别手机号、邮箱、身份证、银行卡、账号、密码、token、secret 等类型，并输出掩码后的样例值。访问审批模块提供申请、通过、驳回和有效期控制。"),
        ("3.2.5 审计日志与统计看板模块", "审计模块记录操作用户、操作类型、目标类型、目标 id、时间、IP、结果和详情。统计看板读取资产数量、敏感字段数量、等级/类别统计和最近审计日志。"),
        ("3.2.6 数据安全治理 Agent 模块", "Agent 模块包含字段分类、审批建议、审计风险、安全报告和问答服务。它综合字段名、样例值、分类规则、申请理由和审计日志输出建议、置信度、告警和整改方案。"),
    ]
    for title, body in modules:
        add_heading(doc, title, 3)
        add_para(doc, body)
    add_table(
        doc,
        ["前端页面", "接口", "功能"],
        [
            ["AgentWorkbench.vue", "GET /api/agent/workbench", "展示今日任务、未处理告警、近期建议和告警。"],
            ["AgentFieldClassify.vue", "POST /api/agent/classify-field/{id}", "分析或应用字段分类建议。"],
            ["AgentApprovalAdvice.vue", "POST /api/agent/review-access-request/{id}", "生成访问审批建议。"],
            ["AgentRiskAlerts.vue", "POST /api/agent/analyze-audit-logs", "分析审计日志并生成风险告警。"],
            ["AgentSecurityReport.vue", "GET /api/agent/security-report", "生成 Markdown 安全治理报告。"],
            ["FloatingAgent.vue / AgentChat.vue", "POST /api/agent/chat", "提供悬浮助手和独立问答页面。"],
        ],
    )

    add_heading(doc, "4 系统测试")
    add_heading(doc, "4.1 测试环境", 2)
    add_para(doc, "测试环境使用 Windows 本地开发环境，后端通过 Maven 启动 Spring Boot 服务，默认连接 MySQL 数据库并加载 schema.sql 与 data.sql；前端通过 npm run dev 启动 Vite 服务；浏览器访问 http://localhost:5173。也可直接执行 sql/mysql_schema_init.sql 初始化数据库。")
    add_heading(doc, "4.2 测试内容与结果分析", 2)
    add_table(
        doc,
        ["编号", "测试项", "测试步骤", "预期结果", "结果"],
        [
            ["T1", "登录认证", "使用 admin/123456 登录", "返回 token 并进入首页", "通过"],
            ["T2", "数据资产 CRUD", "新增数据源、数据表、字段并刷新列表", "列表显示新增记录，数据库写入成功", "通过"],
            ["T3", "自动分类", "启用 phone/id_card/password 等规则后执行自动分类", "字段写入对应类别和 L3/L4/L5 等级", "通过"],
            ["T4", "脱敏展示", "查看 phone、email、id_card、bank_card、password_hash 样例值", "返回格式化脱敏结果", "通过"],
            ["T5", "高敏访问控制", "普通用户访问 L4/L5 原始值且无审批", "返回 requiresApproval=true，仅显示脱敏值并写审计", "通过"],
            ["T6", "访问申请审批", "用户提交申请，审批人通过，再访问原始值", "有效期内可查看原始样例值，审批记录完整", "通过"],
            ["T7", "审计日志", "执行登录、查看字段、审批、分类修改", "audit_log 中存在操作类型、用户、目标和结果", "通过"],
            ["T8", "Agent 字段分类", "对 id_card_no 或 password_hash 执行 Agent 分析", "返回分类、等级、理由和置信度，可选择应用", "通过"],
            ["T9", "Agent 风险分析", "执行审计风险分析", "生成 risk_alert 并可标记处理", "通过"],
            ["T10", "安全报告", "点击生成安全报告", "返回资产概览、分类覆盖率、风险和整改建议", "通过"],
        ],
    )
    add_heading(doc, "4.3 系统效果展示", 2)
    for item in [
        "首页看板展示数据源、数据表、字段、敏感字段数量以及等级/类别统计。",
        "字段资产页面展示字段详情、原始值访问状态和脱敏样例值。",
        "分类分级页面支持人工分类和自动分类，分类结果可按字段、类别、等级查看。",
        "访问申请页面支持普通用户提交申请，审批页面支持审批人通过或驳回。",
        "审计日志页面按时间倒序展示关键操作，便于追溯。",
        "Agent 工作台集中展示分类建议、审批建议、风险告警、安全报告和问答能力。",
    ]:
        add_bullet(doc, item)
    add_para(doc, "由于本报告由源码和脚本生成，系统截图可在正式汇报前按实际运行界面补充到本节。当前正文已覆盖测试环境、测试步骤、预期结果与结果分析。")

    add_heading(doc, "5 总结与展望")
    add_heading(doc, "5.1 项目工作总结", 2)
    add_heading(doc, "5.1.1 完成的主要工作", 3)
    for item in [
        "完成数据分类分级保护系统的需求分析、总体架构、核心业务流程和数据库建模。",
        "设计 20 张核心表，覆盖用户权限、数据资产、分类分级、访问审批、审计日志和 Agent 治理。",
        "完成 MySQL 初始化脚本，包含核心表、约束、索引和演示数据。",
        "实现 Spring Boot 后端 REST API 与 Vue 3 前端页面，支撑完整课程演示流程。",
        "实现本地规则型 Agent，提高字段分类、审批建议、风险分析和报告生成的智能化程度。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.1.2 计划与实际的差异分析", 3)
    add_table(
        doc,
        ["阶段", "计划工作", "实际完成情况", "差异分析"],
        [
            ["需求分析", "完成数据资产、分类分级、审批、审计需求", "按计划完成，并增加 Agent 治理需求", "实际范围略大于计划。"],
            ["数据库设计", "完成 E-R、关系模式和 SQL 建表", "增加视图、存储过程、触发器和索引", "数据库对象更完整。"],
            ["后端实现", "实现基础 CRUD 和审批流程", "完成 CRUD、认证、脱敏、审计、Agent 接口", "功能深度超过计划。"],
            ["前端实现", "实现管理页面和演示页面", "完成看板、资产、分类、审批、审计、Agent 页面", "基本按计划完成。"],
            ["测试汇报", "完成测试用例和汇报材料", "完成测试表、报告和演示流程说明", "截图与视频可在答辩前补充。"],
        ],
    )
    add_heading(doc, "5.1.3 个人工作总结（按成员）", 3)
    add_table(
        doc,
        ["成员", "实际工作内容", "个人总结"],
        [
            ["组长（请填写）", "需求分析、总体架构、数据库模型、报告统稿", "负责把业务安全需求转化为数据库对象和系统流程，保证整体一致性。"],
            ["成员A（请填写）", "后端认证、资产接口、审批、审计、Agent 服务", "完成安全逻辑在服务层和数据库层的落地。"],
            ["成员B（请填写）", "前端页面、路由、表单、看板、Agent 交互", "完成前端业务流程串联，使系统具备可演示界面。"],
            ["成员C（请填写）", "SQL 脚本、测试、文档、图表和汇报材料", "整理数据库对象、测试用例和报告图表，保证提交材料完整。"],
        ],
    )
    add_heading(doc, "5.2 存在的问题与不足", 2)
    for item in [
        "当前演示环境密码为明文初始化，生产环境应使用 BCrypt 或 Argon2 等强哈希算法。",
        "本地规则 Agent 可解释、可离线，但对复杂语义字段的判断能力弱于专业模型。",
        "权限控制主要体现在接口和业务逻辑层，后续可进一步引入数据库行级安全或细粒度策略。",
        "系统截图、录屏和真实小组信息需要在正式提交前补充完善。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 后续工作展望", 2)
    for item in [
        "接入真实数据库元数据扫描能力，自动发现数据源、表结构和字段样例。",
        "引入更细粒度的策略引擎，支持按角色、部门、时间、申请理由和数据等级动态授权。",
        "完善密码加密、接口鉴权、日志防篡改和审计报表导出能力。",
        "结合外部大模型或私有化模型提升 Agent 对复杂业务字段和异常行为的识别能力。",
        "增加数据血缘、风险趋势图、整改闭环工单和合规检查清单，使系统更接近实际企业数据安全治理平台。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "参考文献")
    refs = [
        "全国人民代表大会常务委员会．中华人民共和国网络安全法［Z］．北京：中国法治出版社，2024．",
        "全国人民代表大会常务委员会．中华人民共和国数据安全法［Z］．北京：中国法治出版社，2024．",
        "全国人民代表大会常务委员会．中华人民共和国个人信息保护法［Z］．北京：中国法治出版社，2024．",
        "GB/T 22239-2019．信息安全技术 网络安全等级保护基本要求［S］．北京：中国标准出版社，2019．",
        "GB/T 35273-2020．信息安全技术 个人信息安全规范［S］．北京：中国标准出版社，2020．",
        "GB/T 41479-2022．信息安全技术 网络数据处理安全要求［S］．北京：中国标准出版社，2022．",
        "全国信息安全标准化技术委员会．GB/T 44226-2024 信息安全技术 数据分类分级实施指南［S］．北京：中国标准出版社，2024．",
        "文继荣，李铭轩．AIGC时代网络生成内容合规治理与检测技术框架［J］．北京理工大学学报（社会科学版），2023，25(6)：121-130．",
        "屠要峰，牛家浩，王德政．面向开放大数据环境的动态数据保护系统［J］．软件学报，2023，34(3)：1213-1235．",
        "项目源码与数据库脚本：D:/database/backend，D:/database/frontend，D:/database/sql/mysql_schema_init.sql．",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", False, after=3)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "数据库系统及安全课程设计报告 - 基于分类分级的数据安全治理系统"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            set_run(run, size=9, color=GRAY)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("四川大学网络空间安全学院")
        set_run(run, size=9, color=GRAY)

    props = doc.core_properties
    props.title = "基于分类分级的数据安全治理系统课程设计报告"
    props.author = "课程设计小组"
    props.subject = "数据库系统及安全课程设计"
    props.keywords = "数据安全, 数据库设计, 分类分级, 动态脱敏, 审计日志, Agent"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
